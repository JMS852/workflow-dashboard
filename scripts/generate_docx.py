"""Generate Word document from USER_GUIDE content — 四号宋体."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup ──────────────────────────────────────────────
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

# ── Style helpers ───────────────────────────────────────────

FONT_NAME = '宋体'
FONT_SIZE = Pt(14)  # 四号
FONT_SIZE_H1 = Pt(18)
FONT_SIZE_H2 = Pt(16)
FONT_SIZE_CODE = Pt(11)
FONT_SIZE_TABLE = Pt(12)

def set_font(run, name=FONT_NAME, size=FONT_SIZE, bold=False, color=None):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = color

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, '黑体', FONT_SIZE_H1, True)

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(12)
    run = p.add_run(text)
    set_font(run, size=Pt(12), color=RGBColor(0x8B, 0x8F, 0xA3))

def add_h2(text):
    p = doc.add_paragraph()
    p.space_before = Pt(24)
    p.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, '黑体', FONT_SIZE_H2, True)

def add_h3(text):
    p = doc.add_paragraph()
    p.space_before = Pt(16)
    p.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, bold=True)

def add_para(text, indent=False):
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_font(run)

def add_code(text):
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, 'Consolas', FONT_SIZE_CODE)

def add_note(text):
    p = doc.add_paragraph()
    p.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=Pt(12))

def add_table(headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, '黑体', FONT_SIZE_TABLE, True)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=FONT_SIZE_TABLE)
    doc.add_paragraph()  # spacing

# ══════════════════════════════════════════════════════════

add_title('Workflow Dashboard — 使用说明书')
add_subtitle('多 Agent 协作信差平台 · v2.0\n以 Claude Code / Codex 为 Agent 驱动复杂工程任务')

# ── 1 ─────────────────────────────────────────────────────
add_h2('1. 核心概念')

add_h3('Dashboard 是什么？')
add_para('Workflow Dashboard 是一个信差平台。它不产生 AI 内容——只负责在用户和多个 Claude Code / Codex Agent 之间传递消息、汇聚结论并推动辩论和决策。', indent=True)

add_h3('与传统 AI 工具的根本区别')
add_table(
    ['', '传统 AI 工具 (ChatGPT 等)', 'Workflow Dashboard'],
    [
        ['AI 来源', '单一 API 调用', '多个独立 Claude Code / Codex CLI 实例'],
        ['工作方式', '一问一答', '多 Agent 并行产出 → 互审辩论 → 投票决策'],
        ['Dashboard 角色', 'AI 执行者（自己调 API）', '信差——传递消息、汇聚结论'],
        ['适用任务', '简单对话', '复杂工程：算法、数学、架构、嵌入式'],
        ['Token 效率', 'N/A', 'Agent 间只交换结论 ~200 字，节省 25×+'],
        ['上下文', '单次对话', '每 Agent 保持独立完整上下文，跨轮次记忆'],
    ]
)

add_h3('三段式工作流')
add_code('Round 1 产出  → 用户提任务 → 所有 Agent 并行产出方案 + 200 字结论')
add_code('Round 2 辩论  → Dashboard 汇聚结论对比表 → Agent 互相评判、改进')
add_code('Round 3 决策  → Agent 投票 → Dashboard 展示最终决策')

add_note('📌 三条设计原则：\n'
         '1. Dashboard 不产生 AI 内容 — 只做消息传递和展示\n'
         '2. 每个 Agent 在自己的 Session 里完整工作 — 有上下文、有记忆\n'
         '3. Agent 间只交换结论（~200 字），不互读完整产出 — 大量节省 Token')

# ── 2 ─────────────────────────────────────────────────────
add_h2('2. 安装与启动')

add_h3('系统要求')
add_table(
    ['项目', '要求'],
    [
        ['操作系统', 'Windows 11 / 10（macOS / Linux 理论兼容）'],
        ['Node.js', '18+（推荐 20 LTS）'],
        ['Python', '3.12（仅可选 Sandbox 功能需要）'],
        ['磁盘', '~500 MB（含 Node 依赖）'],
        ['网络', 'Agent 运行时需 Claude API / OpenAI API 连接'],
    ]
)

add_h3('安装 Claude Code CLI')
add_code('npm install -g @anthropic-ai/claude-code')
add_code('# 或: winget install Anthropic.ClaudeCode')
add_code('claude --version    # 验证安装')
add_code('claude login        # 认证（或设置 ANTHROPIC_API_KEY 环境变量）')

add_h3('安装 Codex CLI (OpenAI)')
add_code('npm install -g @openai/codex')
add_code('codex --version     # 验证安装')
add_code('# 认证: set OPENAI_API_KEY=sk-...')

add_h3('安装 Dashboard')
add_code('git clone https://github.com/JMS852/workflow-dashboard.git')
add_code('cd workflow-dashboard')
add_code('npm install')

add_h3('启动')
add_code('node launcher.js')
add_para('启动后自动完成：编译 Electron TS → 构建 React 前端 → 检测 CLI → 打开窗口。关闭窗口即缩到系统托盘。', indent=True)

add_h3('启动故障速查')
add_note('❌ Cannot find module \'electron\' → npm install 未完成，清理重装: rm -rf node_modules && npm install')
add_note('❌ 窗口一闪而过 → TypeScript 编译失败，手动检查: npx tsc -p tsconfig.electron.json')

# ── 3 ─────────────────────────────────────────────────────
add_h2('3. 首次使用')

add_para('步骤 1：打开项目目录。点击「选择项目目录」，Dashboard 在目录下创建 .multi-ai-workflow/ 存放工作记录。', indent=True)

add_para('步骤 2：添加 Agent。右侧 Agent 面板点 + 按钮，填写 ID（如 cc-arch）、类型（Claude Code 或 Codex）、显示名（如 "Claude #1（系统架构师）"）、工作目录。点「注册」，绿色 ON 标签表示已启用。', indent=True)

add_para('步骤 3：提交任务。在输入框输入任务描述，Enter 发送（Shift+Enter 换行）。', indent=True)

add_para('步骤 4：观察流程。状态栏自动显示 Round 1 → 2 → 3 → 完成。切换到「结论对比」和「辩论&决策」标签查看结果。', indent=True)

# ── 4 ─────────────────────────────────────────────────────
add_h2('4. 管理 Agent')

add_h3('Agent 状态指示')
add_table(
    ['图标', '状态', '含义'],
    [
        ['●', 'idle', '空闲中'],
        ['◉ (脉冲)', 'working', '正在执行'],
        ['✓', 'done', '本轮完成'],
        ['✗', 'error', '出错（CLI 崩溃/超时/异常）'],
        ['⚠', '不可用', 'CLI 未找到或路径错误'],
    ]
)

add_h3('推荐配置')
add_table(
    ['ID', '类型', '显示名', '适合场景'],
    [
        ['cc-arch', 'Claude Code', 'Claude #1（系统架构师）', '架构设计、模块划分'],
        ['cc-algo', 'Claude Code', 'Claude #2（算法专家）', '算法、数学建模'],
        ['cc-code', 'Claude Code', 'Claude #3（代码实现）', '代码生成、测试'],
        ['cdx-review', 'Codex', 'Codex（代码审查员）', '代码审查、安全审计'],
    ]
)

add_note('💡 最少 3 个 Agent 才能形成有效对抗验证。2 个输出 + 1 个审查者是最低配置。')

# ── 5 ─────────────────────────────────────────────────────
add_h2('5. 工作流详解')

add_h3('完整数据流')
add_code('用户任务')
add_code('  ├─→ Agent A (--session-id UUID-a) → [完整产出] → 提取结论A')
add_code('  ├─→ Agent B (--session-id UUID-b) → [完整产出] → 提取结论B')
add_code('  └─→ Agent C (--session-id UUID-c) → [完整产出] → 提取结论C')
add_code('      │')
add_code('      ▼ 所有结论到齐 → 自动触发 Round 2')
add_code('      │')
add_code(' 结论对比表: | A | 结论A | | B | 结论B | | C | 结论C |')
add_code('      │')
add_code('  ├─→ Agent A (--resume UUID-a) → 读结论表 → 评判各方')
add_code('  ├─→ Agent B (--resume UUID-b) → 读结论表 → 评判各方')
add_code('  └─→ Agent C (--resume UUID-c) → 读结论表 → 评判各方')
add_code('      │')
add_code('      ▼ 所有辩论到齐 → 自动触发 Round 3 → 投票 → 🏆 最终决策')

add_h3('状态机')
add_table(
    ['状态', '触发条件', '自动动作'],
    [
        ['idle', '初始', '等待用户'],
        ['round_1_produce', '用户提交任务', '并行启动所有 Agent'],
        ['round_2_debate', '全部 R1 完成', '构建结论表 → 并行辩论'],
        ['round_3_decide', '全部 R2 完成', '构建辩论汇总 → 并行决策'],
        ['complete', '全部 R3 完成', '展示最终决策'],
    ]
)

add_note('⏱ 每轮自动推进，无需手动操作。用户可随时点「取消」中断工作流。')

add_h3('查看完整产出')
add_para('「对话」标签中，每个 Agent 消息下有「查看完整产出」折叠区（最多 5000 字）。完整 session 记录保存在 .multi-ai-workflow/sessions/{agentId}.json。', indent=True)

# ── 6 ─────────────────────────────────────────────────────
add_h2('6. Agent 协议与结论格式')

add_para('Agent 之间只读结论，不读完整产出。这要求每个 Agent 的结论必须简洁（200 字以内）、准确（包含核心方案、关键决策、注意点）、机器可解析（通过 ──結論── 标记提取）。', indent=True)

add_h3('结论标记')
add_code('──結論──')
add_code('<结论文本，200 字以内>')
add_code('────────')

add_para('如果 Agent 输出中没有这个标记，Dashboard 会回退取最后 500 字作为结论。协议指令会被自动追加到每个 prompt 末尾。', indent=True)

# ── 7 ─────────────────────────────────────────────────────
add_h2('7. Session 持久化与跨轮次上下文')

add_h3('跨轮次上下文保留机制')
add_table(
    ['CLI', 'Round 1', 'Round 2/3', '存储位置'],
    [
        ['Claude Code', '--session-id <UUID>', '--resume <UUID>', '~/.claude/'],
        ['Codex', 'codex exec（自动生成 ID）', 'codex exec resume <ID>', '~/.codex/'],
    ]
)

add_h3('已验证：上下文确实保留')
add_para('实测结果（双摆模拟测试，2 Agent × 3 轮）：', indent=True)
add_para('• Round 2 中 cc-1 精确引用了自己 Round 1 的三项混沌验证方案', indent=True)
add_para('• Round 2 中 cc-2 精确引用了 cc-1 的拉格朗日推导 + RK45 方案', indent=True)
add_para('• Round 3 中 cc-2 明确说"文件权限问题而非技术失败，代码已就绪可运行"——它记得 Round 1 被权限卡住', indent=True)
add_para('• Round 3 决策综合了三个轮次所有 Agent 的观点', indent=True)

add_note('⚠ Claude Code 原生 session 有 ~30 天闲置后可能清理。Dashboard 的 session store (.multi-ai-workflow/sessions/) 是永久保存的。')

add_h3('崩溃恢复')
add_para('重启 Dashboard → 自动读取 .multi-ai-workflow/sessions/ → 恢复 nativeSessionId → 下次调用自动 --resume → 完整上下文恢复。', indent=True)

# ── 8 ─────────────────────────────────────────────────────
add_h2('8. 使用场景')

add_h3('场景一：算法设计')
add_code('设计一个支持 O(log n) 范围查询和更新的数据结构，')
add_code('用于股票价格区间最值查询。TypeScript 实现 + 复杂度证明。')
add_para('建议 Agent：1 算法专家 + 1 数据结构专家 + 1 代码审查者', indent=True)

add_h3('场景二：数学建模')
add_code('建立传染病 SIR 模型，模拟疫苗接种率变化下的传播。')
add_code('推导微分方程、参数敏感性分析、生成可视化的 Python 脚本。')
add_para('建议 Agent：1 数学建模专家 + 1 数值仿真专家 + 1 验证者', indent=True)

add_h3('场景三：代码审查')
add_code('审查 src/engine/ 目录的 Python 代码，')
add_code('找出性能问题、安全隐患、非 Pythonic 写法。')
add_para('建议 Agent：2 代码审查者 + 1 安全审查者', indent=True)

add_h3('场景四：嵌入式 / 电子竞赛')
add_code('STM32 PWM 电机控制：PID 自整定、编码器反馈、串口调试。')
add_code('C 语言 + HAL 库实现。')
add_para('建议 Agent：1 嵌入式专家 + 1 控制算法专家 + 1 代码实现者', indent=True)

add_h3('场景五：系统架构')
add_code('设计 10 万 QPS 实时消息推送系统。')
add_code('架构图、技术选型、数据流、容错方案。')
add_para('建议 Agent：1 架构师 + 1 网络专家 + 1 可靠性工程师', indent=True)

# ── 9 ─────────────────────────────────────────────────────
add_h2('9. 常见问题与排查')

add_h3('启动问题')
add_note('❌ Cannot find module \'electron\'\n'
         '原因：npm install 未完成。解决：rm -rf node_modules && npm install')
add_note('❌ 窗口一闪而过\n'
         '原因：TypeScript 编译失败。解决：npx tsc -p tsconfig.electron.json 手动检查')

add_h3('Agent 问题')
add_note('❌ Agent 旁出现 ⚠ 图标（CLI 未找到）\n'
         'PATH 中找不到 CLI。解决方法：\n'
         '1. 确认 CLI 已安装：claude --version / codex --version\n'
         '2. 安装到 PATH 可见位置：npm install -g @anthropic-ai/claude-code\n'
         '3. 或手动指定完整路径。\n'
         'Windows 常见路径：\n'
         '  Claude: %LOCALAPPDATA%\\Microsoft\\WinGet\\Packages\\Anthropic.ClaudeCode_*\\claude.exe\n'
         '  Codex:  %LOCALAPPDATA%\\OpenAI\\Codex\\bin\\{hash}\\codex.exe')

add_note('❌ Agent 执行卡住不动\n'
         '可能原因：API 密钥未设 → echo %ANTHROPIC_API_KEY%\n'
         '网络问题 → 手动测试 claude -p "hello"\n'
         'CLI 进程卡死 → 点「取消」终止\n'
         '超时（5 分钟）→ 任务太复杂，拆分或增加超时')

add_note('❌ Agent 返回 exit code 1\n'
         '常见错误及解决：\n'
         '• Invalid session ID → Dashboard 自动生成 UUID\n'
         '• --resume requires valid session → 确保 Round 1 先成功\n'
         '• Authentication required → claude login 或设置环境变量\n'
         '• Permission denied → Dashboard 已内置 --permission-mode bypassPermissions')

add_note('❌ 结论没有被正确提取\n'
         'Agent 输出中缺少 ──結論── 标记。Dashboard 回退取最后 500 字。协议指令已在 prompt 中，大多数 Agent 会遵守。')

add_note('❌ 一个 Agent 拖慢全局\n'
         '流程等待所有 Agent 完成。解决方法：关闭慢的 Agent（ON→OFF）或点「取消」重试。')

add_note('❌ Codex Round 2 恢复失败\n'
         '可能原因：session ID 提取失败或用 --ephemeral 标志。排查：codex resume --last --include-non-interactive')

add_note('✅ Dashboard 崩溃后 Agent 上下文还在吗？\n'
         '还在。重启后自动读 sessions 文件 → 恢复 nativeSessionId → 下次调用自动 --resume。')

add_note('✅ Token 消耗有多省钱？\n'
         '传统方式（全文交换）：~30,000 token 用于"理解别人"\n'
         '结论模式：~1,200 token\n'
         '节省：96%（25×）')

# ── 10 ────────────────────────────────────────────────────
add_h2('10. 高级配置')

add_h3('环境变量')
add_table(
    ['变量', '用途'],
    [
        ['ANTHROPIC_API_KEY', 'Claude API 密钥'],
        ['OPENAI_API_KEY', 'OpenAI API 密钥（Codex）'],
        ['CLAUDE_CODE_PATH', '手动指定 Claude Code CLI 完整路径'],
        ['CODEX_PATH', '手动指定 Codex CLI 完整路径'],
    ]
)

add_h3('CLI 自动发现顺序')
add_para('1. 环境变量 CLAUDE_CODE_PATH / CODEX_PATH', indent=True)
add_para('2. PATH 中的 claude / codex', indent=True)
add_para('3. %LOCALAPPDATA%\\Microsoft\\WinGet\\Packages\\（winget 安装）', indent=True)
add_para('4. %LOCALAPPDATA%\\OpenAI\\Codex\\bin\\（Codex 桌面应用）', indent=True)
add_para('5. %APPDATA%\\npm\\（npm 全局安装）', indent=True)

add_h3('自定义模型')
add_para('在 agent-manager.ts 的 buildCliArgs 中添加 --model 参数可为不同 Agent 指定模型容量：', indent=True)
add_code('// 快 Agent（低成本）')
add_code('args.push(\'--model\', \'claude-haiku-4-5\');')
add_code('// 强 Agent（高质量）')
add_code('args.push(\'--model\', \'claude-sonnet-4-6\');')

add_h3('MQTT 外部任务源（可选）')
add_para('保留的 engine/mqtt_client.py 可从外部 MQTT Broker 接收任务发布到 workflow/tasks/#，适用于 CI/CD 集成。', indent=True)

add_h3('Sandbox（可选）')
add_para('保留的 engine/sandbox.py 可在 Docker/子进程中执行 Agent 产出的 Python 代码，验证正确性。', indent=True)

# ── Footer ─────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Workflow Dashboard v2.0 — 信差平台\n基于 MQTT-3388 事件驱动架构\nhttps://github.com/JMS852/workflow-dashboard')
set_font(run, size=Pt(10), color=RGBColor(0x8B, 0x8F, 0xA3))

# ── Save ───────────────────────────────────────────────────
output_path = os.path.expanduser('~/Desktop/Workflow_Dashboard_使用说明书_v2.0.docx')
doc.save(output_path)
print(f'Saved to: {output_path}')
